import docx, os

def init(links, variables):
    ex_variables = []
    for var in variables['word']:
        ex_variables.append(var)

    for var in variables['default']:
        ex_variables.append(var)
    
    for link in links:
        """ Для документов word 2013+ """
        if link.count('docx'):
            doc = docx.Document(link)
            findReplace(doc, ex_variables)

            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.tables:
                                for table in cell.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            findReplace(cell, ex_variables)
                                            doc.save(link)
                            else:
                                findReplace(cell, ex_variables)
                                doc.save(link)

            doc.save(link)



def findReplace(obj, variables):
    """ Производит поиск и замену в документах."""
    for paragraph in obj.paragraphs:
        for var in variables:
            if paragraph.text.count(var['var']):
                print(var['var'], end='\n\n')
                paragraph.text = paragraph.text.replace(var['var'], str(var['val']))