from docx.shared import RGBColor
pronouns = [
    'олжен',
    'олжна',
    'олжны',
    'олжно',
]

layout = [{
        'олжен': {'ым': 'ый', 'им': 'ий', 'ый': 'ый', 'ий': 'ий'},
        'олжна': {'ой': 'ая', 'ая': 'ая'},
        'олжны': {'ыми': 'ые', 'ими': 'ие', 'ые': 'ые'},
        'олжно': {'ым': 'ое', 'им': 'ие', 'ое': 'ое'}
    },{
        'олжен': {
            'влять': 'вляет', 'вовать': 'вует',
            'яться': 'яется', 'овать': 'ует',
            'чать': 'чает', 'еть': 'еет', 'ять': 'ит',
            'ить': 'ит', 'ать': 'ает'
        },
        'олжна': {
            'влять': 'вляет', 'вовать': 'вует',
            'яться': 'яется', 'овать': 'ует',
            'чать': 'чает', 'еть': 'еет', 'ять': 'ит',
            'ить': 'ит', 'ать': 'ает'
        },
        'олжны': {
            'влять': 'вляют', 'вовать': 'вуют',
            'яться': 'яются', 'овать': 'уют',
            'чать': 'чают', 'еть': 'еют', 'ять': 'ят',
            'ить': 'ят', 'ать': 'ают'
        },
        'олжно': {
            'влять': 'вляет', 'вовать': 'вует',
            'яться': 'яется', 'овать': 'ует',
            'чать': 'чает', 'еть': 'еет', 'ять': 'ит',
            'ить': 'ит', 'ать': 'ает'
        }
    }
]
endings = [
    'ый', 'ой', 'ыми', 'ими', 'ым', 'им', 'ая', 'ые', 'ое', 'ий'
], [
    'влять', 'вовать', 'яться', 'овать', 'еть', 'ять', 'ить', 'чать', 'ать'
]

def handling(doc):
    """ Ищет в документе совпадение с pronouns. 
    
        При успехе создает [список] из 2, 3 или 4 слов
        далее вызывает _get_an_ending([список])
    """
    _start(doc)
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        for table in cell.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    _start(cell)
                    else:
                        _start(cell)

def _start(obj):
    for paragraph in obj.paragraphs:
        for index, run in enumerate(paragraph.runs):
            for pronoun in pronouns:
                if (run.text.count(pronoun)): # если есть совпадение
                    temp = []
                    w1 = paragraph.runs[index].text
                    temp.append(w1)
                    w2 = paragraph.runs[index + 1].text
                    temp.append(w2)
                    if w2.count('быть'):
                        w3 = paragraph.runs[index + 2].text
                        temp.append(w3)
                        result = _get_an_ending(temp)
                        if result:
                            set_word(index, paragraph.runs, result, 2)
                        else:
                            try:
                                w4 = paragraph.runs[index + 3].text
                            except IndexError:
                                print('Output or range')
                                break
                            temp.append(w4)
                            result = _get_an_ending(temp)
                            if result:
                                set_word(index, paragraph.runs, result, 3)
                            else:
                                _set_warning(index, paragraph.runs, 3)
                    else:
                        result = _get_an_ending(temp)
                        if result:
                            set_word(index, paragraph.runs, result, 1)
                        else:
                            _set_warning(index, paragraph.runs, 1)

def _get_an_ending(words):
    """ Возвращает обработаенную строку."""

    if len(words) == 3:
        cont = 2
    elif len(words) == 4:
        cont = 3
    if words[1].count('быть'):
        for pronoun in pronouns:
            if words[0].count(pronoun):
                for end in endings[0]:
                    if words[cont].count(end):
                        try:
                            set_ending = layout[0][pronoun][end]
                        except KeyError:
                            print("no math endings in", words[cont])
                            return False
                        if cont == 2:
                            result = words[2].replace(end, set_ending)
                            return result
                        elif cont == 3:
                            for end2 in endings[0]:
                                if words[2].count(end2):
                                    second_ending = layout[0][pronoun][end2]
                                    result = '%s ' % words[2].replace(end2,                              second_ending)
                                    result += words[3].replace(end, set_ending)

                            result = words[2] + words[3].replace(end,                                        set_ending)
                            return result

    elif len(words) == 2:
        word = words[1]
        for pronoun in pronouns:
            if words[0].count(pronoun):
                for end in endings[1]:
                    if word.count(end):
                        set_eding = layout[1][pronoun][end]
                        result = word.replace(end, set_eding)
                        return result

def set_word(index, runs, result, count):
    color = RGBColor(0x33, 0x99, 0x55)
    if count == 1:
        runs[index].clear()
        runs[index + 1].text = result
        runs[index + 1].italic = True
        runs[index + 1].font.color.rgb = color
    elif count == 2:
        runs[index].clear()
        runs[index + 1].clear()
        runs[index + 2].text = result
        runs[index + 2].italic = True
        runs[index + 2].font.color.rgb = color
    elif count == 3:
        runs[index].clear()
        runs[index + 1].clear()
        runs[index + 2].clear()
        runs[index + 3].text = result
        runs[index + 3].italic = True
        runs[index + 3].font.color.rgb = color

def _set_warning(index, runs, count):
    color = RGBColor(0xff, 0x65, 0x0b)
    if count == 1:
        runs[index].font.color.rgb = color
        runs[index + 1].font.color.rgb = color

    elif count == 2:
        runs[index].font.color.rgb = color
        runs[index + 1].font.color.rgb = color
        runs[index + 2].font.color.rgb = color

    elif count == 3:
        runs[index].font.color.rgb = color
        runs[index + 1].font.color.rgb = color
        runs[index + 2].font.color.rgb = color
        runs[index + 3].font.color.rgb = color

# res = _get_an_ending(['должно ', 'стираться'])

# print(res)