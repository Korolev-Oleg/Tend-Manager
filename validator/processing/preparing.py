import docx

def init(doc):
    """ Разбирает весь документ на Runs."""
    # for paragraph in doc.paragraphs:
    #     temp = []
    #     for run in paragraph.runs:
    #         words = run.text.split(' ')
    #         for word in words:
    #             temp.append(word)
    #     paragraph.clear()
    #     for word in temp:
    #         paragraph.add_run('%s ' % word)
    _split(doc)
    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        for table in cell.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    _split(cell)
                    else:
                        _split(cell)
       
        

def _split(obj):
    for paragraph in obj.paragraphs:
        temp = []
        for run in paragraph.runs:
            words = run.text.split(' ')
            for word in words:
                temp.append(word)
        paragraph.clear()
        for word in temp:
            paragraph.add_run('%s ' % word)