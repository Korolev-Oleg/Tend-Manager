import docx, os

# file = "test.docx"
# doc = docx.Document(file)
""" list methods.
    
    Keyword arguments:
        'add_heading', 'add_page_break', 'add_paragraph', 'add_picture', 'add_section', 'add_table', 'core_properties', 'element', 'inline_shapes', 'paragraphs', 'part', 'save', 'sections', 'settings', 'styles', 'tables'
"""

# doc.tables
""" [...]. """

# doc.tables[0]
def doc_tables():
    """ .
        
        Keyword arguments:
            'add_column', 'add_row', 'alignment', 'autofit', 'cell', 'column_cells', 'columns', 'part', 'row_cells', 'rows', 'style', 'table', 'table_direction'

            doc.tables -> rows -> cells -> paragrphs...
    """
    pass


# doc.paragraphs
""" [...]. """

def doc_paragraphs():
    """ list methods.
        
        Keyword arguments:
            'add_run', 'alignment', 'clear', 'insert_paragraph_before', 'paragraph_format', 'part', 'runs', 'style', 'text'
    """
    # doc.paragraphs[0]
    pass

# doc.paragraphs[0].runs
""" .
    
    Keyword arguments:
    
"""

# doc.paragraphs[0].runs[0]
""" .
    
    Keyword arguments:
        'add_break', 'add_picture', 'add_tab', 'add_text', 'bold', 'clear', 'element', 'font', 'italic', 'part', 'style', 'text', 'underline'
"""




doc = docx.Document('Заявка на участие.docx')

# for table in doc.tables:
#             for row in table.rows:
#                 for cn, cell in enumerate(row.cells):
#                     for pn, paragraph in enumerate(cell.paragraphs):
#                         for rn, run in enumerate(paragraph.runs):
#                             print(cn, pn, rn, run.text, end='\n\n')

splt = doc.tables[0].rows[3].cells[2].paragraphs[1].runs[0].text.split(' ')

print(splt)